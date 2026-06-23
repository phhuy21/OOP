#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Generate BAO_CAO_SKYGATE.docx — SkyGate Airport Management System report.
Format matches BAO_CAO_DO_AN_1.docx (Times New Roman 13pt, navy/blue headings).

Cải tiến:
 - Code & cây thư mục đặt trong khung viền + nền xám.
 - Danh mục Hình/Bảng có link bấm được (bookmark + PAGEREF) như Mục lục.
 - Nội dung khớp code thực tế: SQLite WAL, repo GitHub đúng, thông báo lỗi thật.
"""

from docx import Document
import os
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

d = Document()

# ============ PAGE / STYLE SETUP (match reference) ============
sec = d.sections[0]
sec.page_width  = Emu(7772400)
sec.page_height = Emu(10058400)
sec.left_margin   = Emu(1080135)
sec.right_margin  = Emu(720090)
sec.top_margin    = Emu(899795)
sec.bottom_margin = Emu(899795)

style_normal = d.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(13)
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.15
_rpr = style_normal.element.get_or_add_rPr()
_rfonts = OxmlElement('w:rFonts')
_rfonts.set(qn('w:eastAsia'), 'Times New Roman')
_rpr.append(_rfonts)

for hname, sz, col in [('Heading 1', 14, (0x36, 0x5F, 0x91)),
                       ('Heading 2', 13, (0x4F, 0x81, 0xBD)),
                       ('Heading 3', 13, (0x4F, 0x81, 0xBD))]:
    st = d.styles[hname]
    st.font.name = 'Times New Roman'
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = RGBColor(*col)

CENTER = WD_ALIGN_PARAGRAPH.CENTER

# ============ LOW-LEVEL XML HELPERS ============
_bm_id = [1000]

def _shade_and_border(paragraph, fill='F4F4F4', border='B0B0B0'):
    """Thêm nền màu + viền 4 cạnh cho 1 paragraph (khung code/cây thư mục)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '6')
        e.set(qn('w:space'), '6'); e.set(qn('w:color'), border)
        pBdr.append(e)
    pPr.append(pBdr)

def _bookmark(paragraph, name, text, italic=False, bold=False):
    """Đặt bookmark quanh text trong paragraph (đích cho link danh mục)."""
    _bm_id[0] += 1
    bid = str(_bm_id[0])
    bs = OxmlElement('w:bookmarkStart'); bs.set(qn('w:id'), bid); bs.set(qn('w:name'), name)
    paragraph._p.append(bs)
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'; run.italic = italic; run.bold = bold
    be = OxmlElement('w:bookmarkEnd'); be.set(qn('w:id'), bid)
    paragraph._p.append(be)
    return run

def _list_entry(name, text, indent=0.0):
    """Một dòng danh mục có link bấm được + số trang (PAGEREF) + dot leader."""
    p = d.add_paragraph()
    pf = p.paragraph_format
    if indent: pf.left_indent = Cm(indent)
    pf.space_after = Pt(3)
    pf.tab_stops.add_tab_stop(Cm(16.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    hyper = OxmlElement('w:hyperlink'); hyper.set(qn('w:anchor'), name)
    # run: text
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'), 'Times New Roman'); rf.set(qn('w:hAnsi'), 'Times New Roman')
    rpr.append(rf)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '26'); rpr.append(sz)
    r.append(rpr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); hyper.append(r)
    # tab
    rt = OxmlElement('w:r'); rt.append(OxmlElement('w:tab')); hyper.append(rt)
    # PAGEREF field
    for kind, val in [('begin', None), ('instr', f' PAGEREF {name} \\h '),
                      ('separate', None), ('text', '1'), ('end', None)]:
        rr = OxmlElement('w:r')
        if kind == 'instr':
            it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = val; rr.append(it)
        elif kind == 'text':
            tt = OxmlElement('w:t'); tt.text = val; rr.append(tt)
        else:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), kind); rr.append(fc)
        hyper.append(rr)
    p._p.append(hyper)
    return p

# ============ CONTENT HELPERS ============
def para(text, bold=False, align=None, sa=None, sb=None, italic=False):
    p = d.add_paragraph()
    run = p.add_run(text); run.font.name = 'Times New Roman'
    if bold: run.bold = True
    if italic: run.italic = True
    if align is not None: p.alignment = align
    if sa is not None: p.paragraph_format.space_after = sa
    if sb is not None: p.paragraph_format.space_before = sb
    return p

def h1(text): d.add_heading(text, level=1)
def h2(text): d.add_heading(text, level=2)
def h3(text): d.add_heading(text, level=3)

def bullet(text):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(" - " + text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    return p

def code_block(code_text):
    """Đoạn code trong khung viền + nền xám, font Consolas."""
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3); p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.keep_together = True
    _shade_and_border(p, fill='F4F4F4')
    lines = code_text.split('\n')
    for i, line in enumerate(lines):
        run = p.add_run(line); run.font.name = 'Consolas'; run.font.size = Pt(10)
        if i < len(lines) - 1: run.add_break()
    return p

def tree_block(text):
    """Cây thư mục trong khung viền + nền xanh nhạt."""
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3); p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.keep_together = True
    _shade_and_border(p, fill='EAF1FB', border='9CB8DC')
    lines = text.split('\n')
    for i, line in enumerate(lines):
        run = p.add_run(line); run.font.name = 'Consolas'; run.font.size = Pt(10)
        if i < len(lines) - 1: run.add_break()
    return p

def _shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def table(headers, rows):
    t = d.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, head in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = ''
        _shade_cell(cell, 'D9E2F3')  # xanh nhạt cho header
        run = cell.paragraphs[0].add_run(head)
        run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]; cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    d.add_paragraph()
    return t

# ---- Figure / Table captions (đích link cho danh mục) ----
FIG = {
    'f31': 'Hình 3.1 — Sơ đồ lớp UML (Class Diagram)',
    'f32': 'Hình 3.2 — Sơ đồ trạng thái chuyến bay (State Diagram)',
    'f33': 'Hình 3.3 — Kiến trúc tổng thể hệ thống (mô hình 3 tầng)',
    'f41': 'Hình 4.1 — Cấu trúc thư mục mã nguồn SkyGate',
    'f51': 'Hình 5.1 — Giao diện đăng nhập',
    'f52': 'Hình 5.2 — Dashboard tổng quan',
    'f53': 'Hình 5.3 — Quản lý chuyến bay',
    'f54': 'Hình 5.4 — Đặt vé và chọn ghế trực quan',
    'f55': 'Hình 5.5 — Check-in và cảnh báo hành lý',
    'f56': 'Hình 5.6 — Phân quyền theo vai trò',
}
TAB = {
    't11': 'Bảng 1.1 — Công nghệ sử dụng trong dự án',
    't31': 'Bảng 3.1 — Yêu cầu chức năng của hệ thống',
    't32': 'Bảng 3.2 — Các lớp nhóm Con người (people/)',
    't33': 'Bảng 3.3 — Các lớp nhóm Máy bay (aircraft/)',
    't34': 'Bảng 3.4 — Các lớp nhóm Vận hành (operations/)',
    't35': 'Bảng 3.5 — Các lớp nhóm Hệ thống (system/, auth/, common/, web/)',
    't36': 'Bảng 3.6 — Lược đồ các bảng trong CSDL SQLite',
    't41': 'Bảng 4.1 — Các quy tắc nghiệp vụ và xử lý vi phạm',
    't42': 'Bảng 4.2 — Các endpoint chính của REST API',
}

def fig_caption(key):
    p = d.add_paragraph(); p.alignment = CENTER
    _bookmark(p, key, FIG[key], italic=True)
    return p

def tab_caption(key):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    _bookmark(p, key, FIG.get(key) or TAB[key], bold=True)
    return p

# ============================================================
# COVER PAGE
# ============================================================
para('BỘ GIÁO DỤC VÀ ĐÀO TẠO', bold=True, align=CENTER, sa=Pt(2))
para('TRƯỜNG ĐẠI HỌC CÔNG NGHỆ KỸ THUẬT TP. HCM', bold=True, align=CENTER, sa=Pt(2))
para('KHOA ĐIỆN – ĐIỆN TỬ', bold=True, align=CENTER, sa=Pt(2))
para('🙢🙝🕮🙝🙢', align=CENTER, sa=Pt(20))
para('BÁO CÁO BÀI TẬP LỚN', bold=True, align=CENTER, sa=Pt(12))
para('HỆ THỐNG QUẢN LÝ SÂN BAY SKYGATE', bold=True, align=CENTER, sa=Pt(6))
para('(SkyGate Airport Management System)', italic=True, align=CENTER, sa=Pt(20))
para('Môn: Lập trình Hướng đối tượng', align=CENTER, sa=Pt(4))
para('GVHD: ThS. Trần Thị Quỳnh Như', bold=True, align=CENTER, sa=Pt(20))
para('Nhóm 5 — Thành viên:', bold=True, align=CENTER, sa=Pt(8))
for name, mssv in [('Nguyễn Ngọc Trường Phi', '23119187'),
                   ('Vũ Minh Quang', '23119198'),
                   ('Nguyễn Thị Thúy Hằng', '23119142'),
                   ('Nguyễn Đức Lộc', '23119165'),
                   ('Phạm Gia Huy', '23119148')]:
    para(f'{name} – {mssv}', align=CENTER, sa=Pt(2))
para('', sa=Pt(28))
para('Thành phố Hồ Chí Minh, Tháng 6 năm 2026', align=CENTER)

# ============================================================
# LỜI CẢM ƠN
# ============================================================
d.add_page_break()
h1('LỜI CẢM ƠN')
para('Đầu tiên, nhóm chúng em xin gửi lời cảm ơn chân thành và sâu sắc nhất đến giảng viên '
     'hướng dẫn — ThS. Trần Thị Quỳnh Như. Cô đã tận tình định hướng, góp ý và chỉ ra những '
     'điểm cần cải thiện trong suốt quá trình nhóm thực hiện bài tập lớn này. Những buổi trao '
     'đổi với cô không chỉ giúp nhóm hoàn thiện sản phẩm mà còn củng cố thêm hiểu biết về tư '
     'duy lập trình hướng đối tượng — nền tảng cốt lõi của ngành học.')
para('Nhóm cũng xin cảm ơn các thành viên trong nhóm 5 đã nỗ lực phối hợp, đóng góp ý tưởng '
     'và san sẻ công việc trong một đề tài có phạm vi rộng và độ phức tạp cao. Tinh thần làm '
     'việc nhóm và trách nhiệm của từng người là yếu tố then chốt để bài tập lớn hoàn thành '
     'đúng tiến độ và đạt chất lượng mong muốn.')
para('Trong quá trình thực hiện, do hạn chế về thời gian và kinh nghiệm, báo cáo chắc chắn '
     'không tránh khỏi thiếu sót. Nhóm rất mong nhận được những ý kiến đóng góp quý báu từ cô '
     'để sản phẩm ngày càng hoàn thiện hơn.')
para('Nhóm 5 xin trân trọng cảm ơn.', sb=Pt(12))

# ============================================================
# MỤC LỤC
# ============================================================
d.add_page_break()
h1('MỤC LỤC')
p = d.add_paragraph()
for kind, val in [('begin', None), ('instr', ' TOC \\o "1-3" \\h \\z \\u '), ('separate', None)]:
    r = p.add_run()
    if kind == 'instr':
        it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = val; r._r.append(it)
    else:
        fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), kind); r._r.append(fc)
ri = p.add_run('(Nhấn Ctrl+A rồi F9 để cập nhật mục lục, danh mục và số trang)')
ri.italic = True; ri.font.size = Pt(11)
rend = p.add_run(); fce = OxmlElement('w:fldChar'); fce.set(qn('w:fldCharType'), 'end'); rend._r.append(fce)

# ============================================================
# DANH MỤC BẢNG  (link bấm được)
# ============================================================
d.add_page_break()
h1('DANH MỤC BẢNG')
for k in ['t11', 't31', 't32', 't33', 't34', 't35', 't36', 't41', 't42']:
    _list_entry(k, TAB[k])

# ============================================================
# DANH MỤC HÌNH  (link bấm được)
# ============================================================
d.add_page_break()
h1('DANH MỤC HÌNH')
for k in ['f31', 'f32', 'f33', 'f41', 'f51', 'f52', 'f53', 'f54', 'f55', 'f56']:
    _list_entry(k, FIG[k])

# ============================================================
# DANH MỤC TỪ VIẾT TẮT
# ============================================================
d.add_page_break()
h1('DANH MỤC TỪ VIẾT TẮT')
table(['Viết tắt', 'Đầy đủ'], [
    ['API', 'Application Programming Interface'],
    ['CLI', 'Command Line Interface'],
    ['CRUD', 'Create, Read, Update, Delete'],
    ['CSDL', 'Cơ sở dữ liệu'],
    ['CSS', 'Cascading Style Sheets'],
    ['HTML', 'HyperText Markup Language'],
    ['HTTP', 'HyperText Transfer Protocol'],
    ['JSON', 'JavaScript Object Notation'],
    ['OOP', 'Object-Oriented Programming (Lập trình hướng đối tượng)'],
    ['REST', 'Representational State Transfer'],
    ['SKG', 'SkyGate (mã sân bay nhà)'],
    ['SQL', 'Structured Query Language'],
    ['STL', 'Standard Template Library (C++)'],
    ['UML', 'Unified Modeling Language'],
    ['WAL', 'Write-Ahead Logging (chế độ ghi của SQLite)'],
])

# ============================================================
# CHƯƠNG 1
# ============================================================
d.add_page_break()
h1('CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI')

h2('1.1. Lý do chọn đề tài')
for t in [
    'Sân bay là một hệ thống phức tạp trong thực tế, gồm nhiều đối tượng tương tác: chuyến '
    'bay, máy bay, phi công, hành khách, cổng, đường băng, hành lý... Đây là bài toán lý tưởng '
    'để áp dụng lập trình hướng đối tượng (OOP) vì mỗi thực thể có thể mô hình hóa thành một '
    'lớp với thuộc tính và hành vi riêng.',
    'Nhu cầu quản lý sân bay rất thực tế: theo dõi trạng thái chuyến bay, quản lý hành khách '
    'từ khâu mua vé đến khi lên máy bay, điều phối tài nguyên (gate, đường băng, máy bay) sao '
    'cho không xung đột, và cảnh báo các tình huống bất thường như hành lý quá cân hay chuyến '
    'bay bị hoãn.',
    'Bài toán có tính nghiệp vụ rõ ràng, dễ mô hình hóa bằng các lớp và quan hệ kế thừa — phù '
    'hợp với mục tiêu rèn luyện tư duy thiết kế lớp và áp dụng đóng gói, kế thừa, đa hình, '
    'trừu tượng.',
    'Nhóm mong muốn xây dựng một sản phẩm hoàn chỉnh vượt ra ngoài phạm vi bài tập console đơn '
    'thuần, bao gồm cả giao diện Web để có trải nghiệm trực quan và thực tế hơn.',
]:
    para(t)

h2('1.2. Mục tiêu đề tài')
para('Nhóm xác định các mục tiêu chính của đề tài như sau:', sa=Pt(6))
for t in [
    'Luyện tập phân tích bài toán thực tế theo hướng đối tượng: xác định lớp, thuộc tính, '
    'phương thức và quan hệ.',
    'Thiết kế hệ thống lớp có phân cấp kế thừa rõ ràng, tận dụng đa hình để xử lý khác biệt '
    'giữa các loại máy bay và vai trò người dùng.',
    'Áp dụng đóng gói để bảo vệ dữ liệu và đảm bảo trạng thái hợp lệ của từng đối tượng.',
    'Xây dựng chương trình mô phỏng quản lý sân bay hoàn chỉnh với hai giao diện: console (CLI) '
    'và web (REST API + HTML/CSS/JS).',
    'Xử lý các tình huống nghiệp vụ phức tạp: xung đột lịch trình, kiểm tra chứng chỉ phi công, '
    'giới hạn giờ bay/thời gian nghỉ, tương thích máy bay–gate–đường băng, vòng đời chuyến bay '
    'qua máy trạng thái.',
    'Triển khai hệ thống phân quyền 3 vai trò (Admin/Staff/Customer) với chức năng khác nhau.',
]:
    bullet(t)

h2('1.3. Phạm vi và giới hạn')
h3('1.3.1. Trong phạm vi')
for t in [
    'Quản lý 1 sân bay nhà (SKG — SkyGate) với đầy đủ cổng (gate) và đường băng (runway).',
    'Quản lý chuyến bay theo vòng đời 12 trạng thái (10 trạng thái chính + Delayed/Cancelled).',
    'Quản lý 3 loại máy bay: thân rộng (WideBody), thân hẹp (NarrowBody), cánh quạt (Turboprop).',
    'Quản lý hành khách: thêm, đặt chỗ, check-in, boarding, theo dõi trạng thái (kể cả No-show).',
    'Quản lý hành lý: cảnh báo khi vượt mức tiêu chuẩn (>2 kiện hoặc >46 kg tổng), chặn cứng khi '
    'vượt giới hạn tối đa (>3 kiện hoặc >50 kg).',
    'Hệ thống đặt/huỷ vé (Ticket) cho Customer với sơ đồ chọn ghế, phân hạng vé (Thương gia/Thường).',
    'Phân quyền 3 vai trò: Admin (toàn quyền), Staff (vận hành), Customer (tra cứu + mua vé).',
    'Lưu trữ dữ liệu bằng cơ sở dữ liệu SQLite (chế độ WAL) cho phép CLI và Web truy cập đồng thời.',
    'Mô phỏng đồng hồ ảo (tickSimulation), hàng đợi cất/hạ cánh, mô phỏng thời tiết xấu.',
]:
    bullet(t)

h3('1.3.2. Ngoài phạm vi')
for t in [
    'Không triển khai thanh toán thực tế hay tích hợp cổng thanh toán.',
    'Không mô phỏng nhiều sân bay đồng thời (chỉ quản lý sân bay nhà SKG; các điểm đến là tên '
    'thành phố, không phải sân bay trong hệ thống).',
    'Không mô phỏng thời tiết theo thời gian thực (chỉ mô phỏng thời tiết xấu thủ công).',
    'Không có bản đồ trực quan (phiên bản Leaflet map đã được gỡ bỏ khỏi giao diện chính).',
    'Tên "SkyGate" và tên model máy bay (SkyCruiser-900, AeroSwift-321...) là tên giả lập.',
]:
    bullet(t)

h2('1.4. Công nghệ sử dụng')
tab_caption('t11')
table(['Thành phần', 'Công nghệ'], [
    ['Ngôn ngữ chính', 'C++17 (g++ 15.x — MSYS2 UCRT64)'],
    ['Giao diện CLI', 'Console (src/main.cpp) — nhập/xuất văn bản thuần'],
    ['Web Server', 'cpp-httplib — C++ single-header HTTP server'],
    ['Frontend Web', 'HTML5 + CSS3 + JavaScript (Vanilla, giao diện Dark Theme, không framework)'],
    ['Xử lý JSON', 'JSON builder thủ công (src/web/Json.h) — không dùng thư viện ngoài'],
    ['Lưu trữ dữ liệu', 'SQLite (amalgamation) — chế độ WAL, file data/skygate.db'],
    ['Quản lý mã nguồn', 'Git / GitHub'],
    ['Công cụ biên dịch', 'bash build.sh / build_web.sh (biên dịch trực tiếp g++, không CMake)'],
])

# ============================================================
# CHƯƠNG 2
# ============================================================
d.add_page_break()
h1('CHƯƠNG 2: CƠ SỞ LÝ THUYẾT')

h2('2.1. Giới thiệu lập trình hướng đối tượng (OOP)')
para('Lập trình hướng đối tượng (Object-Oriented Programming — OOP) là phương pháp lập trình '
     'lấy "đối tượng" (object) làm trung tâm. Mỗi đối tượng kết hợp dữ liệu (thuộc tính) và '
     'hành vi (phương thức). OOP ra đời từ những năm 1960 với Simula 67 và phát triển mạnh qua '
     'Smalltalk, C++, Java, C#, Python...')
para('Khác với lập trình hướng thủ tục — tổ chức chương trình thành các hàm xử lý dữ liệu '
     'tách rời — OOP tổ chức chương trình thành các lớp (class). Mỗi lớp đóng gói dữ liệu và '
     'phương thức thao tác trên dữ liệu đó, giúp mã nguồn dễ bảo trì, dễ mở rộng và gần gũi với '
     'cách con người nhận thức thế giới thực.')
para('Trong bài tập lớn này, OOP là nền tảng để mô hình hóa toàn bộ hệ thống sân bay: mỗi thực '
     'thể là một lớp, các quan hệ được thể hiện qua kế thừa, thành phần (composition) và liên kết.')

h2('2.2. Bốn tính chất cơ bản của OOP')

h3('2.2.1. Tính đóng gói (Encapsulation)')
para('Đóng gói là cơ chế che giấu dữ liệu nội bộ, chỉ truy cập qua phương thức công khai, bảo '
     'vệ tính toàn vẹn dữ liệu.')
para('Áp dụng trong SkyGate:', bold=True)
for t in [
    'Thuộc tính quan trọng khai báo private/protected: trạng thái chuyến bay (status_), số giờ '
    'bay phi công (monthlyFlightHours_)...',
    'Truy cập qua getter/setter có kiểm tra: Person::setAge() chỉ nhận tuổi không âm; '
    'Pilot::addFlightHours() không cho cộng giờ âm.',
    'Lớp Flight che giấu status_ và passengers_ — mọi chuyển trạng thái phải qua advance().',
    'Lớp Baggage che giấu giới hạn mặc định (2 kiện / 46 kg) và chỉ công khai qua phương thức kiểm tra.',
]:
    bullet(t)

h3('2.2.2. Tính kế thừa (Inheritance)')
para('Kế thừa cho phép lớp con kế thừa thuộc tính/phương thức từ lớp cha, thể hiện quan hệ "is-a".')
para('Áp dụng trong SkyGate:', bold=True)
for t in [
    'Con người: Person (abstract) → Passenger, Staff; Staff (abstract) → Pilot, GroundStaff.',
    'Máy bay: Aircraft (abstract) → WideBodyAircraft, NarrowBodyAircraft, TurbopropAircraft.',
    'Người dùng: auth::User (abstract) → Admin, Staff, Customer.',
]:
    bullet(t)

h3('2.2.3. Tính đa hình (Polymorphism)')
para('Đa hình cho phép cùng một phương thức có hành vi khác nhau tùy đối tượng — qua hàm ảo và '
     'con trỏ/tham chiếu lớp cha.')
para('Áp dụng trong SkyGate:', bold=True)
for t in [
    'requiredRunwayLength() — WideBody trả 2800m, NarrowBody 2500m, Turboprop 1500m; gọi trên '
    'con trỏ Aircraft* cho kết quả tùy loại máy bay thực tế.',
    'role() — Passenger trả "Passenger", Pilot trả "Pilot"...; mỗi đối tượng tự khai báo vai trò.',
    'menu() trong auth::User — Admin/Staff/Customer trả danh sách tab khác nhau, giao diện web '
    'dùng chung một vòng lặp để dựng tab.',
]:
    bullet(t)

h3('2.2.4. Tính trừu tượng (Abstraction)')
para('Trừu tượng là mô hình hóa khái niệm thực tế thành lớp/đối tượng, ẩn chi tiết triển khai. '
     'Lớp trừu tượng chứa hàm ảo thuần và không thể khởi tạo trực tiếp.')
para('Áp dụng trong SkyGate:', bold=True)
for t in [
    'Aircraft là lớp trừu tượng định nghĩa giao diện chung cho mọi loại máy bay — phải tạo qua '
    'lớp con hoặc AircraftFactory.',
    'Person là nền tảng chung cho mọi con người trong hệ thống.',
    'auth::User định nghĩa giao diện xác thực và phân quyền.',
]:
    bullet(t)

h2('2.3. Các khái niệm OOP nâng cao đã sử dụng')
for t in [
    'Composition (Thành phần): Flight chứa Aircraft, Crew, Gate, danh sách Passenger; Airport '
    'chứa Gate và Runway; Crew chứa Pilot và GroundStaff.',
    'Factory Pattern: AircraftFactory::create(category, ...) tạo đúng lớp con máy bay; '
    'auth::makeUser(role, ...) tạo đúng lớp con tài khoản.',
    'Enum class: FlightStatus, AircraftCategory, GateType, StaffRole — type-safe.',
    'Smart Pointers: std::shared_ptr quản lý bộ nhớ tự động cho mọi đối tượng.',
    'State Machine: Flight quản lý vòng đời 12 trạng thái với điều kiện chuyển trạng thái.',
    'OpResult pattern: mọi thao tác nghiệp vụ trả về OpResult{ok, message} thay vì throw exception.',
]:
    bullet(t)

h2('2.4. Giới thiệu ngôn ngữ C++ và các thư viện sử dụng')
para('C++ được chọn vì là ngôn ngữ chính của môn học, hỗ trợ đầy đủ OOP và cho phép kiểm soát '
     'bộ nhớ ở mức thấp. Dự án dùng chuẩn C++17.')
para('Các thư viện và công cụ chính:', bold=True, sb=Pt(6))
for t in [
    'STL: std::vector, std::map, std::set, std::shared_ptr, std::function.',
    'cpp-httplib (third_party/httplib.h): HTTP server single-header — routing GET/POST, phục vụ '
    'file tĩnh từ web/.',
    'SQLite (third_party/sqlite3.c): cơ sở dữ liệu nhúng, dùng ở chế độ WAL cho truy cập đồng thời.',
    'JSON thủ công (src/web/Json.h): lớp Obj và Arr xây chuỗi JSON — không phụ thuộc thư viện ngoài.',
    'std::filesystem: kiểm tra/tạo thư mục data/, đổi tên file khi di trú dữ liệu.',
]:
    bullet(t)

# ============================================================
# CHƯƠNG 3
# ============================================================
d.add_page_break()
h1('CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG')

h2('3.1. Phân tích yêu cầu')
h3('3.1.1. Yêu cầu chức năng')
tab_caption('t31')
table(['STT', 'Chức năng', 'Mô tả'], [
    ['1', 'Quản lý sân bay', 'CRUD sân bay, cổng (Gate), đường băng (Runway)'],
    ['2', 'Quản lý máy bay', 'Thêm/xem/xoá máy bay theo loại (WideBody, NarrowBody, Turboprop)'],
    ['3', 'Quản lý chuyến bay', 'Tạo, cập nhật trạng thái theo state machine, gán máy bay/gate'],
    ['4', 'Quản lý hành khách', 'Đặt chỗ, check-in, boarding, theo dõi trạng thái (kể cả No-show)'],
    ['5', 'Quản lý hành lý', 'Cảnh báo quá mức (>2 kiện hoặc >46 kg); chặn cứng (>3 kiện hoặc >50 kg)'],
    ['6', 'Điều phối sân bay', 'Hàng đợi cất/hạ cánh, ưu tiên khẩn cấp, mô phỏng thời tiết xấu'],
    ['7', 'Phân quyền', 'Admin / Staff / Customer — kiểm tra server-side ở mọi route mutation'],
    ['8', 'Đặt/Huỷ vé', 'Customer mua vé → tạo Passenger + Ticket, gán ghế; huỷ vé → trả ghế'],
    ['9', 'Lưu/đọc dữ liệu', 'CSDL SQLite (WAL); tự động di trú từ định dạng .txt cũ'],
    ['10', 'Đồng hồ mô phỏng', 'Tua thời gian ảo, tự động đẩy trạng thái chuyến theo lịch'],
    ['11', 'Giao diện Web', 'Dashboard tab theo vai trò, giao tiếp REST API'],
])
h3('3.1.2. Yêu cầu phi chức năng')
for t in [
    'Giao diện trực quan, dễ sử dụng — cả CLI (menu phân cấp) và Web (tab theo vai trò).',
    'Dữ liệu nhất quán — mọi thao tác ghi đều qua AirportSystem để kiểm tra ràng buộc.',
    'Đồng bộ giữa CLI và Web — cùng đọc/ghi một CSDL SQLite ở chế độ WAL.',
    'Xử lý lỗi rõ ràng — mọi thao tác trả về OpResult{ok, message} bằng tiếng Việt.',
]:
    bullet(t)

h2('3.2. Xác định các đối tượng (Classes)')
h3('Nhóm 1: Con người — src/people/')
tab_caption('t32')
table(['Lớp', 'Vai trò', 'Kế thừa từ', 'Tính chất OOP'], [
    ['Person', 'Lớp cha trừu tượng — id, name, age, phone', '—', 'Trừu tượng, Đóng gói'],
    ['Passenger', 'Hành khách — passport, flight, seat, checkedIn, boarded, Baggage', 'Person', 'Kế thừa, Đa hình'],
    ['Staff', 'Nhân viên sân bay (trừu tượng) — baseAirport', 'Person', 'Kế thừa, Trừu tượng'],
    ['Pilot', 'Phi công — chứng chỉ, giờ bay ≤100h, nghỉ ≥8h', 'Staff', 'Kế thừa, Đóng gói'],
    ['GroundStaff', 'Nhân viên mặt đất — department', 'Staff', 'Kế thừa, Đa hình'],
])
h3('Nhóm 2: Máy bay — src/aircraft/')
tab_caption('t33')
table(['Lớp', 'Vai trò', 'Kế thừa từ', 'Tính chất OOP'], [
    ['Aircraft', 'Lớp cha trừu tượng — registration, model, capacity + 5 hàm ảo thuần', '—', 'Trừu tượng'],
    ['WideBodyAircraft', 'Thân rộng — 2800m RW, gate đôi, quay đầu 0 phút (đã bỏ)', 'Aircraft', 'Kế thừa, Đa hình'],
    ['NarrowBodyAircraft', 'Thân hẹp — 2500m RW, gate đơn, quay đầu 0 phút (đã bỏ)', 'Aircraft', 'Kế thừa, Đa hình'],
    ['TurbopropAircraft', 'Cánh quạt — 1500m RW, gate bãi ngoài, quay đầu 0 phút (đã bỏ)', 'Aircraft', 'Kế thừa, Đa hình'],
    ['AircraftFactory', 'Factory Pattern — tạo máy bay đúng lớp con theo AircraftCategory', '—', 'Factory Pattern'],
])
h3('Nhóm 3: Vận hành — src/operations/')
tab_caption('t34')
table(['Lớp', 'Vai trò', 'Tính chất OOP'], [
    ['Airport', 'Sân bay — chứa Gate, Runway; tìm gate trống phù hợp', 'Composition'],
    ['Flight', 'Chuyến bay — state machine, chứa Aircraft/Crew/Gate/Passenger', 'Composition, State Machine'],
    ['Gate', 'Cổng — loại gate + danh sách Reservation, chống xung đột', 'Đóng gói'],
    ['Runway', 'Đường băng — mã + chiều dài (mét)', 'Đóng gói'],
    ['Crew', 'Tổ bay — Pilot + GroundStaff, validate cho chuyến bay', 'Composition'],
    ['Baggage', 'Hành lý — số kiện + khối lượng, kiểm tra quá mức', 'Đóng gói'],
    ['Ticket', 'Vé — mã vé, chuyến, hành khách, chủ vé, hạng ghế', 'Đóng gói'],
])
h3('Nhóm 4: Hệ thống — src/system/, auth/, common/, web/')
tab_caption('t35')
table(['Lớp', 'Vai trò', 'Tính chất OOP'], [
    ['AirportSystem', 'Lớp điều khiển trung tâm (~1000 dòng) — CRUD, nghiệp vụ, lưu/đọc CSDL, mô phỏng', 'Controller'],
    ['Database', 'Wrapper SQLite — WAL mode, prepared statement, transaction', 'Utility / RAII'],
    ['OpResult', 'Kết quả thao tác {ok, message} — không dùng exception cho business rule', 'Value Object'],
    ['DateTime', 'Thời điểm (YYYY-MM-DD HH:MM) — so sánh, cộng/trừ phút, kiểm tra giao khoảng', 'Value Object'],
    ['Enums', 'FlightStatus, AircraftCategory, GateType, StaffRole + toString/fromString', 'Enum Class'],
    ['auth::User', 'Lớp cha trừu tượng — Admin/Staff/Customer', 'Trừu tượng, Đa hình'],
    ['json::Obj / Arr', 'JSON builder thủ công cho Web API', 'Utility'],
])

h2('3.3. Sơ đồ lớp (Class Diagram)')
para('Sơ đồ lớp UML dưới đây thể hiện các quan hệ kế thừa, thành phần và liên kết giữa các lớp '
     'chính trong hệ thống SkyGate (cú pháp Mermaid).', bold=True, sa=Pt(6))
code_block(
    "classDiagram\n"
    "    Person <|-- Passenger\n"
    "    Person <|-- Staff\n"
    "    Staff <|-- Pilot\n"
    "    Staff <|-- GroundStaff\n"
    "    Aircraft <|-- WideBodyAircraft\n"
    "    Aircraft <|-- NarrowBodyAircraft\n"
    "    Aircraft <|-- TurbopropAircraft\n"
    "    User <|-- Admin\n"
    "    User <|-- StaffAuth\n"
    "    User <|-- Customer\n"
    "    Flight *-- Aircraft\n"
    "    Flight *-- Crew\n"
    "    Flight o-- Passenger\n"
    "    Crew o-- Pilot\n"
    "    Crew o-- GroundStaff\n"
    "    Airport *-- Gate\n"
    "    Airport *-- Runway\n"
    "    Passenger *-- Baggage\n"
    "    AirportSystem *-- Airport\n"
    "    AirportSystem o-- Flight\n"
    "    AirportSystem o-- Aircraft\n"
    "    AirportSystem o-- User\n"
    "    AirportSystem o-- Ticket\n"
    "    AirportSystem *-- Database\n"
    "    class Person {\n"
    "        #string id_\n"
    "        #string fullName_\n"
    "        #int age_\n"
    "        +role()* string\n"
    "    }\n"
    "    class Aircraft {\n"
    "        #string registration_\n"
    "        #int capacity_\n"
    "        +requiredRunwayLength()* int\n"
    "        +minGateRank()* int\n"
    "    }\n"
    "    class Flight {\n"
    "        -FlightStatus status_\n"
    "        -vector~Passenger~ passengers_\n"
    "        +advance() OpResult\n"
    "        +checkInPassenger() OpResult\n"
    "    }"
)
fig_caption('f31')

h2('3.4. Sơ đồ trạng thái chuyến bay (State Diagram)')
para('Flight quản lý vòng đời qua máy trạng thái gồm 12 trạng thái. Mỗi chuyển trạng thái được '
     'kiểm tra điều kiện trước khi thực hiện.', sa=Pt(6))
code_block(
    "stateDiagram-v2\n"
    "    [*] --> Scheduled\n"
    "    Scheduled --> CheckIn\n"
    "    CheckIn --> Boarding\n"
    "    Boarding --> GateClosed\n"
    "    GateClosed --> Ready\n"
    "    Ready --> Takeoff\n"
    "    Takeoff --> InAir\n"
    "    InAir --> Landed\n"
    "    Landed --> Turnaround\n"
    "    Turnaround --> Completed\n"
    "    Completed --> [*]\n"
    "    Scheduled --> Delayed\n"
    "    CheckIn --> Delayed\n"
    "    Boarding --> Delayed\n"
    "    Delayed --> Boarding : tiếp tục sau hoãn\n"
    "    Scheduled --> Cancelled\n"
    "    CheckIn --> Cancelled\n"
    "    Boarding --> Cancelled"
)
fig_caption('f32')
para('Đặc điểm của máy trạng thái:', bold=True)
for t in [
    '10 trạng thái chính tạo vòng đời tuyến tính của một chuyến bay bình thường.',
    '2 trạng thái sự cố: Delayed (hoãn) và Cancelled (huỷ).',
    'Từ Delayed có thể quay lại Boarding khi advance() được gọi.',
    'Không thể hoãn/huỷ khi chuyến đang bay (Takeoff → Landed) hoặc đã kết thúc.',
    'Điều kiện chuyển sang Ready: phải có máy bay + gate; tổ bay là tùy chọn (nếu có thì phải hợp lệ).',
    'Khi hạ cánh (Landed): tự động cộng giờ bay và cập nhật mốc nghỉ cho phi công.',
]:
    bullet(t)

h2('3.5. Kiến trúc hệ thống')
para('SkyGate áp dụng kiến trúc 3 tầng (3-tier):', sa=Pt(6))
for t in [
    'Tầng Frontend (web/): HTML + CSS + JavaScript — dashboard tabbed, gọi Fetch API, tự động '
    'attach actor từ localStorage.',
    'Tầng Backend (src/web/web_main.cpp): cpp-httplib server — nhận request, gọi AirportSystem, '
    'serialize kết quả thành JSON. KHÔNG chứa business logic.',
    'Tầng Data (data/skygate.db): CSDL SQLite chế độ WAL — lưu toàn bộ trạng thái, cho phép CLI '
    'và Web truy cập đồng thời.',
]:
    bullet(t)
para('AirportSystem (src/system/) là lớp trung tâm — cả CLI và Web đều gọi cùng bộ phương thức '
     'public, đảm bảo logic nghiệp vụ nhất quán.', bold=True, sb=Pt(6))
code_block(
    "graph TB\n"
    "    subgraph Frontend\n"
    "        CLI[Console App<br>main.cpp]\n"
    "        Web[Web Browser<br>HTML/CSS/JS]\n"
    "    end\n"
    "    subgraph Backend[C++ Backend]\n"
    "        HTTP[httplib Server<br>web_main.cpp]\n"
    "        AS[AirportSystem<br>Business Logic]\n"
    "        DB[Database<br>SQLite wrapper]\n"
    "    end\n"
    "    Files[(data/skygate.db<br>SQLite WAL)]\n"
    "    CLI -->|gọi trực tiếp| AS\n"
    "    Web -->|HTTP API / JSON| HTTP\n"
    "    HTTP -->|public methods| AS\n"
    "    AS --> DB\n"
    "    DB -->|read/write đồng thời| Files"
)
fig_caption('f33')

h2('3.6. Thiết kế cơ sở dữ liệu (SQLite)')
para('SkyGate lưu toàn bộ trạng thái trong một file CSDL SQLite duy nhất (data/skygate.db), mở '
     'ở chế độ WAL (Write-Ahead Logging). WAL cho phép nhiều tiến trình đọc đồng thời và một '
     'tiến trình ghi mà không khoá lẫn nhau — nhờ đó tiến trình CLI và Web cùng làm việc trên '
     'một CSDL, dữ liệu luôn nhất quán giữa hai giao diện.')
para('Khi mở CSDL, hệ thống đặt các PRAGMA: journal_mode=WAL, foreign_keys=ON, synchronous='
     'NORMAL, busy_timeout=5000ms. Quan hệ nhiều–nhiều (tổ bay ↔ phi công, chuyến bay ↔ hành '
     'khách) được chuẩn hóa thành bảng nối (junction table) thay vì chuỗi nối bằng dấu phẩy.')
para('Di trú dữ liệu: nếu phát hiện định dạng .txt cũ mà chưa có skygate.db, hệ thống tự động '
     'đọc toàn bộ .txt, ghi vào SQLite, rồi đổi tên các file .txt thành .txt.bak (an toàn, có '
     'thể khôi phục).', sa=Pt(6))
tab_caption('t36')
table(['Bảng', 'Vai trò', 'Khóa chính'], [
    ['airports', 'Sân bay', 'code'],
    ['runways', 'Đường băng', '(airport_code, code)'],
    ['gates', 'Cổng ra máy bay', '(airport_code, code)'],
    ['aircraft', 'Máy bay', 'registration'],
    ['pilots', 'Phi công (chứng chỉ, giờ bay, mốc nghỉ)', 'id'],
    ['groundstaff', 'Nhân viên mặt đất', 'id'],
    ['passengers', 'Hành khách (trạng thái, hành lý)', 'id'],
    ['crews', 'Tổ bay', 'id'],
    ['crew_pilots / crew_ground', 'Bảng nối tổ bay ↔ phi công / NV mặt đất', '(crew_id, …_id)'],
    ['flights', 'Chuyến bay (trạng thái, tài nguyên gán)', 'code'],
    ['flight_passengers', 'Bảng nối chuyến bay ↔ hành khách (có cột pos)', '(flight_code, passenger_id)'],
    ['users', 'Tài khoản đăng nhập', 'username'],
    ['tickets', 'Vé đã đặt', 'ticket_id'],
    ['queues', 'Hàng đợi cất/hạ cánh', '(type, pos)'],
    ['metadata', 'Trạng thái hệ thống (ticket_counter, simulation_time)', 'key'],
    ['event_log', 'Nhật ký sự kiện mô phỏng', 'id (AUTOINCREMENT)'],
])
para('Ghi chú: Đặt chỗ gate (Gate::Reservation) không lưu trong CSDL mà được dựng lại trong bộ '
     'nhớ sau khi nạp dữ liệu (rebuildGateReservations), dựa trên quan hệ chuyến bay ↔ gate.',
     italic=True)

h2('3.7. Thiết kế giao diện Web')
para('Giao diện Web xây dựng theo mô hình SPA đơn giản: index.html + app.js + style.css, không '
     'framework, không bước build — chỉnh sửa và reload trực tiếp. Giao diện dùng Dark Theme '
     '(nền tối) hiện đại.', sa=Pt(6))
for t in [
    'Layout: sidebar điều hướng (trái) + vùng nội dung chính (phải); tab hiển thị theo vai trò '
    '(từ menu() của User).',
    'Đăng nhập: overlay form username/password; sau khi login, actor lưu vào localStorage và tự '
    'động attach vào mọi POST.',
    'Admin thấy nhiều tab nhất (gồm Sân bay, Máy bay, Tạo chuyến, Tài khoản...); Staff bị giới '
    'hạn; Customer chỉ thấy tra cứu + đặt vé.',
    'Phân quyền: client-side ẩn/hiện tab; server-side kiểm tra role ở mọi POST — server-side là '
    'gate thật.',
    'Giám sát (renderMonitor): tính client-side từ /api/state — cảnh báo khách chưa check-in gần '
    'giờ bay, chuyến gần đầy/đầy, hành lý quá cân.',
    'Tự động lưu: sau mỗi POST thành công (trừ /api/login), server gọi saveAll("data") để ghi '
    'vào SQLite.',
]:
    bullet(t)

# ============================================================
# CHƯƠNG 4
# ============================================================
d.add_page_break()
h1('CHƯƠNG 4: CÀI ĐẶT VÀ TRIỂN KHAI')

h2('4.1. Cấu trúc thư mục mã nguồn')
tree_block(
    "skygate/\n"
    "├── src/\n"
    "│   ├── aircraft/       # Aircraft (abstract), WideBody, NarrowBody,\n"
    "│   │                   #   Turboprop, AircraftFactory\n"
    "│   ├── auth/           # User (abstract), Admin, Staff, Customer, makeUser()\n"
    "│   ├── common/         # DateTime, Enums, OpResult, Utils, Database (SQLite)\n"
    "│   ├── operations/     # Airport, Flight, Gate, Runway, Crew, Baggage, Ticket\n"
    "│   ├── people/         # Person (abstract), Passenger, Staff, Pilot, GroundStaff\n"
    "│   ├── system/         # AirportSystem (~1000 dòng) — bộ điều khiển trung tâm\n"
    "│   ├── web/            # Json (Obj/Arr), web_main.cpp (REST API)\n"
    "│   └── main.cpp        # Điểm khởi chạy console\n"
    "├── web/                # Frontend: index.html, app.js, style.css (Dark Theme)\n"
    "├── data/               # CSDL: skygate.db (SQLite WAL)\n"
    "├── third_party/        # httplib.h, sqlite3.c, sqlite3.h\n"
    "├── build.sh            # Build console app\n"
    "├── build_web.sh         # Build web server\n"
    "└── README.md"
)
fig_caption('f41')

h2('4.2. Cài đặt các lớp chính (kèm đoạn code tiêu biểu)')
para('Phần này trích dẫn những đoạn code thể hiện rõ nhất các nguyên lý OOP, kèm giải thích.',
     bold=True, sa=Pt(8))

h3('4.2.1. Kế thừa — Hệ thống phân cấp Aircraft')
para('Lớp Aircraft định nghĩa 5 hàm ảo thuần; mỗi lớp con override để cung cấp thông số riêng.', sa=Pt(6))
code_block(
    "// Aircraft.h — lớp cha trừu tượng\n"
    "class Aircraft {\n"
    "public:\n"
    "    virtual ~Aircraft() = default;\n"
    "    virtual AircraftCategory category() const = 0;\n"
    "    virtual int requiredRunwayLength() const = 0;   // mét\n"
    "    virtual int minTurnaroundMinutes() const = 0;   // phút\n"
    "    virtual GateType preferredGateType() const = 0;\n"
    "    virtual int minGateRank() const = 0;\n"
    "    bool canUseRunwayLength(int len) const {\n"
    "        return len >= requiredRunwayLength();\n"
    "    }\n"
    "    bool canUseGateType(GateType t) const {\n"
    "        return gateRank(t) >= minGateRank();\n"
    "    }\n"
    "};\n"
    "\n"
    "// WideBodyAircraft.h — lớp con override mọi hàm ảo thuần\n"
    "class WideBodyAircraft : public Aircraft {\n"
    "public:\n"
    "    AircraftCategory category() const override\n"
    "        { return AircraftCategory::WideBody; }\n"
    "    int requiredRunwayLength() const override { return 2800; }\n"
    "    int minGateRank() const override { return 2; }\n"
    "};"
)
para('Giải thích: Aircraft là lớp cha trừu tượng — không thể khởi tạo trực tiếp. Mỗi lớp con '
     'override hàm ảo thuần để trả thông số kỹ thuật riêng; canUseRunwayLength() định nghĩa một '
     'lần ở lớp cha nhưng hành vi thay đổi tùy lớp con — đó là đa hình.', italic=True)

h3('4.2.2. Đa hình — Kiểm tra điều kiện máy bay')
para('Trong AirportSystem::assignAircraft(), hệ thống gọi requiredRunwayLength() trên con trỏ '
     'Aircraft*; kết quả phụ thuộc loại máy bay thực tế.', sa=Pt(6))
code_block(
    "// Trích AirportSystem::assignAircraft()\n"
    "if (!origin->hasRunwayFor(*ac)) {\n"
    "    return OpResult::failure(\n"
    "        \"Sân bay đi \" + origin->code() +\n"
    "        \" không có đường băng đủ dài (cần >= \" +\n"
    "        std::to_string(ac->requiredRunwayLength()) + \"m).\");\n"
    "}\n"
    "\n"
    "// Airport::hasRunwayFor() — gọi đa hình từ Aircraft\n"
    "bool Airport::hasRunwayFor(const Aircraft& aircraft) const {\n"
    "    for (const auto& r : runways_)\n"
    "        if (aircraft.canUseRunwayLength(r.lengthMeters()))\n"
    "            return true;\n"
    "    return false;\n"
    "}"
)
para('Giải thích: ac có kiểu khai báo shared_ptr<Aircraft> nhưng trỏ tới một lớp con cụ thể. '
     'Lời gọi ac->requiredRunwayLength() được phân giải động (dynamic dispatch) tại runtime.',
     italic=True)

h3('4.2.3. Đóng gói — Quản lý trạng thái chuyến bay')
para('Flight::advance() bảo vệ status_ (private); mọi chuyển trạng thái phải qua phương thức '
     'này để kiểm tra điều kiện.', sa=Pt(6))
code_block(
    "OpResult Flight::advance(const vector<string>& crewIssues) {\n"
    "    if (status_ == FlightStatus::Cancelled)\n"
    "        return OpResult::failure(\"Chuyến bay đã bị huỷ...\");\n"
    "    if (status_ == FlightStatus::Completed)\n"
    "        return OpResult::failure(\"Chuyến bay đã hoàn tất.\");\n"
    "    if (status_ == FlightStatus::Delayed) {\n"
    "        status_ = FlightStatus::Boarding;  // tiếp tục sau hoãn\n"
    "        return OpResult::success(\"... -> Boarding.\");\n"
    "    }\n"
    "    FlightStatus target = nextStatus(status_);\n"
    "    if (target == FlightStatus::Ready) {\n"
    "        if (!aircraft_) return OpResult::failure(\"Chưa gán máy bay.\");\n"
    "        if (gateCode_.empty()) return OpResult::failure(\"Chưa gán gate.\");\n"
    "    }\n"
    "    status_ = target;\n"
    "    if (target == FlightStatus::Landed && crew_)\n"
    "        crew_->applyFlightCompletion(flightHours(), arrival_);\n"
    "    return OpResult::success(\"Trạng thái -> \" + toString(status_));\n"
    "}"
)
para('Giải thích: status_ là private — bên ngoài không thể gán trực tiếp. Mọi chuyển trạng thái '
     'qua advance(), nơi kiểm tra chuyến đã kết thúc chưa, đã có máy bay/gate chưa.', italic=True)

h3('4.2.4. Factory Pattern — AircraftFactory')
para('AircraftFactory đóng gói logic tạo đúng lớp con máy bay theo AircraftCategory.', sa=Pt(6))
code_block(
    "shared_ptr<Aircraft> AircraftFactory::create(\n"
    "        AircraftCategory category, const string& reg,\n"
    "        const string& model, int capacity) {\n"
    "    switch (category) {\n"
    "        case AircraftCategory::WideBody:\n"
    "            return make_shared<WideBodyAircraft>(reg, model, capacity);\n"
    "        case AircraftCategory::NarrowBody:\n"
    "            return make_shared<NarrowBodyAircraft>(reg, model, capacity);\n"
    "        case AircraftCategory::Turboprop:\n"
    "            return make_shared<TurbopropAircraft>(reg, model, capacity);\n"
    "    }\n"
    "    return make_shared<NarrowBodyAircraft>(reg, model, capacity);\n"
    "}"
)
para('Giải thích: Người gọi chỉ truyền category, không cần biết lớp con cụ thể. Thêm loại máy '
     'bay mới chỉ cần thêm case — không sửa code ở nơi sử dụng.', italic=True)

h3('4.2.5. Composition — Flight chứa nhiều đối tượng')
para('Flight "có" máy bay, tổ bay, gate và danh sách hành khách (giữ qua shared_ptr).', sa=Pt(6))
code_block(
    "class Flight {\n"
    "private:\n"
    "    string code_;\n"
    "    FlightStatus status_ = FlightStatus::Scheduled;\n"
    "    shared_ptr<Aircraft> aircraft_;              // composition\n"
    "    shared_ptr<Crew> crew_;                      // composition\n"
    "    string gateCode_;                            // association\n"
    "    vector<shared_ptr<Passenger>> passengers_;   // aggregation\n"
    "};\n"
    "\n"
    "class Crew {\n"
    "    vector<shared_ptr<Pilot>> pilots_;\n"
    "    vector<shared_ptr<GroundStaff>> ground_;\n"
    "};"
)
para('Giải thích: Composition mô hình hóa chính xác cấu trúc thực tế — một chuyến bay cần máy '
     'bay + tổ bay + gate; một tổ bay gồm nhiều phi công và nhân viên.', italic=True)

h2('4.3. Cài đặt các quy tắc nghiệp vụ')
para('Mọi vi phạm đều trả về OpResult::failure() với thông báo tiếng Việt — không throw exception.',
     sa=Pt(6))
tab_caption('t41')
table(['Quy tắc', 'Vị trí kiểm tra', 'Xử lý khi vi phạm'], [
    ['Chứng chỉ máy bay', 'Crew::validateForFlight → Pilot::isCertifiedFor', 'Từ chối gán tổ bay'],
    ['Giới hạn 100 giờ bay/tháng', 'Pilot::canTakeAdditionalHours', 'Từ chối gán tổ bay'],
    ['Nghỉ tối thiểu 8 tiếng', 'Pilot::hasEnoughRestBefore', 'Từ chối gán tổ bay'],
    ['Đường băng đủ dài', 'Airport::hasRunwayFor → Aircraft::canUseRunwayLength', 'Từ chối gán máy bay'],
    ['Gate phù hợp loại máy bay', 'Aircraft::canUseGateType', 'Từ chối gán gate'],
    ['Mặc định cổng duy nhất', 'assignGate (Mặc định khi sân bay có 1 cổng)', 'Tự động gán cổng duy nhất, cho phép dùng chung cổng'],
    ['Xung đột lịch máy bay', 'assignAircraft (kiểm tra trùng giờ bay trực tiếp)', 'Từ chối gán máy bay'],
    ['Hành lý quá mức (>2 kiện / >46 kg)', 'Baggage::isOverweight', 'Cảnh báo (không chặn)'],
    ['Hành lý vượt giới hạn (>3 kiện / >50 kg)', 'setPassengerBaggage / bookTicket', 'Chặn cứng (từ chối)'],
    ['Không hoãn/huỷ chuyến đang bay', 'Flight::delay / cancel', 'Từ chối, thông báo lỗi'],
    ['Boarding khi gate đã đóng', 'Flight::boardPassenger (chỉ khi Boarding)', 'Từ chối, đánh dấu No-show'],
    ['Đặt vé sau giờ bay / đóng quầy', 'bookTicket', 'Từ chối, thông báo lỗi'],
])

h2('4.4. Cài đặt Web Server và REST API')
para('Web server dùng cpp-httplib (single-header). Server phục vụ file tĩnh (web/) và REST API '
     '(/api/*). Mọi logic nghiệp vụ ủy thác cho AirportSystem.', sa=Pt(6))
para('Các endpoint chính:', bold=True)
tab_caption('t42')
table(['Endpoint', 'Method', 'Mô tả', 'Phân quyền'], [
    ['/api/state', 'GET', 'Toàn bộ trạng thái hệ thống (JSON)', 'Tất cả'],
    ['/api/login', 'POST', 'Xác thực username/password', 'Tất cả'],
    ['/api/register', 'POST', 'Đăng ký tài khoản Customer', 'Tất cả'],
    ['/api/flight/create', 'POST', 'Tạo chuyến bay', 'Admin'],
    ['/api/flight/assign-aircraft', 'POST', 'Gán máy bay', 'Admin, Staff'],
    ['/api/flight/assign-gate', 'POST', 'Gán gate (rỗng = tự động)', 'Admin, Staff'],
    ['/api/flight/checkin · board', 'POST', 'Check-in / boarding hành khách', 'Admin, Staff'],
    ['/api/flight/advance', 'POST', 'Chuyển trạng thái kế tiếp', 'Admin, Staff'],
    ['/api/flight/delay · cancel', 'POST', 'Hoãn / huỷ chuyến bay', 'Admin'],
    ['/api/ticket/book · cancel', 'POST', 'Mua / huỷ vé', 'Customer, Staff, Admin'],
    ['/api/aircraft/* · runway/*', 'POST', 'Thêm/xoá máy bay, đường băng', 'Admin'],
    ['/api/sim/tick', 'POST', 'Tua đồng hồ mô phỏng', 'Admin, Staff'],
    ['/api/weather', 'POST', 'Mô phỏng thời tiết xấu', 'Admin'],
    ['/api/users · user/*', 'GET/POST', 'Xem/tạo/xoá tài khoản', 'Admin'],
])
para('Phân quyền kiểm tra server-side bằng lambda hasRole() — đọc actor từ params, tra cứu User, '
     'so sánh role(). Client-side chỉ ẩn/hiện tab.', italic=True)

h2('4.5. Cài đặt giao diện Web (Frontend)')
para('Frontend dùng Vanilla JS — không framework. Ba file chính: index.html, app.js, style.css '
     '(Dark Theme).', sa=Pt(6))
for t in [
    'api(path, params): hàm gọi API chung — GET nếu không params, POST form-encoded nếu có; tự '
    'attach actor từ localStorage.',
    'applyAuth(): hiển thị overlay đăng nhập nếu chưa có actor; sau login gọi /api/state và render.',
    'renderTabs(menu): dựng tab điều hướng từ menu() của User theo vai trò.',
    'renderMonitor(): tính cảnh báo client-side — khách chưa check-in gần giờ bay, chuyến gần '
    'đầy/đầy, hành lý quá cân.',
    'Tự động lưu: sau mỗi POST mutation thành công, server gọi saveAll("data") ghi vào SQLite.',
]:
    bullet(t)

# ============================================================
# CHƯƠNG 5
# ============================================================
d.add_page_break()
h1('CHƯƠNG 5: KẾT QUẢ VÀ KIỂM THỬ')

h2('5.1. Triển khai hệ thống và hướng dẫn chạy chương trình')
h3('Yêu cầu hệ thống')
for t in [
    'Máy chủ: Máy chủ riêng VPS DigitalOcean (Ubuntu Linux) phục vụ chạy Web Server 24/7.',
    'Máy khách: Trình duyệt web hiện đại (Chrome/Edge/Firefox) để truy cập giao diện người dùng.',
    'Công cụ biên dịch: g++ (hỗ trợ C++17) trên Linux VPS hoặc MSYS2 UCRT64 trên Windows.',
]:
    bullet(t)
h3('Triển khai và khởi chạy')
for t in [
    'Triển khai Web Server: Hệ thống được deploy tự động lên VPS qua script deploy.ps1 từ máy trạm.',
    'Lệnh chạy trên VPS: nohup ./skygate_web 8888 > server.log 2>&1 & (biên dịch bằng build_web.sh trên Linux VPS và chạy nền ở cổng 8888).',
    'Địa chỉ truy cập trực tuyến của đề tài: http://159.65.141.209:8888',
    'Ứng dụng Console (CLI): Biên dịch bằng build.sh -> ./skygate.exe để chạy và kiểm tra trực tiếp từ console.',
    'Di trú dữ liệu: Lần đầu chạy, hệ thống tự di trú toàn bộ file .txt cũ sang CSDL SQLite (skygate.db) và lưu trữ dự phòng dưới dạng file .bak.',
    'Tài khoản demo: admin/admin123 (Admin), staff/staff123 (Staff), customer/cus123 (Customer).',
]:
    bullet(t)

h2('5.2. Kết quả chạy thử — Ảnh chụp màn hình')
para('[Phần này cần chèn ảnh chụp màn hình thực tế. Mỗi mục dưới đây mô tả ảnh cần có và vị trí '
     'đặt ảnh; chú thích hình đã được đánh số để liên kết với Danh mục hình.]', italic=True, sa=Pt(8))
shots = [
    ('f51', 'Overlay đăng nhập: ô Username, Password, nút Đăng nhập; thông báo lỗi khi sai mật '
            'khẩu; nút Đăng ký cho Customer mới.'),
    ('f52', 'Dashboard sau khi đăng nhập Admin: đồng hồ mô phỏng, thống kê nhanh (số chuyến bay, '
            'hành khách, máy bay, gate), danh sách chuyến bay kèm trạng thái.'),
    ('f53', 'Tab Chuyến bay: danh sách đầy đủ (mã, điểm đi–đến, giờ, trạng thái, máy bay, gate, '
            'số khách) và các nút Check-in, Board, Advance, Delay, Cancel.'),
    ('f54', 'Tab Đặt vé: chọn chuyến, nhập tên + SĐT, sơ đồ ghế trực quan (2 hàng đầu Thương gia), '
            'khai báo hành lý; hệ thống sinh mã vé TKxxxx.'),
    ('f55', 'Check-in hành khách kèm cảnh báo hành lý vượt mức (VƯỢT MỨC) hiển thị màu nổi bật '
            'trong kết quả và tab Giám sát.'),
    ('f56', 'So sánh 3 giao diện theo vai trò: Admin (nhiều tab nhất), Staff (giới hạn), Customer '
            '(tra cứu + đặt vé). Tab bị ẩn không xuất hiện trong sidebar.'),
]
for key, desc in shots:
    img_name = ""
    if key == 'f51': img_name = 'f51_login.png'
    elif key == 'f52': img_name = 'f52_dashboard.png'
    elif key == 'f53': img_name = 'f52_dashboard.png'
    elif key == 'f54': img_name = 'f57_booking.png'
    elif key == 'f55': img_name = 'f55_monitor.png'
    elif key == 'f56': img_name = 'f56_create.png'
    
    img_path = os.path.join('images', img_name) if img_name else ""
    if img_path and os.path.exists(img_path):
        p = d.add_paragraph()
        p.alignment = CENTER
        p.add_run().add_picture(img_path, width=Cm(15))
    else:
        para('[CHÈN ẢNH CHỤP MÀN HÌNH TẠI ĐÂY]', align=CENTER, sa=Pt(2))
    fig_caption(key)
    para('Mô tả: ' + desc, italic=True, sa=Pt(10))

# ============================================================

# ============================================================
# CHƯƠNG 6
# ============================================================
d.add_page_break()
h1('CHƯƠNG 6: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN')

h2('6.1. Kết quả đạt được')
para('Nhóm 5 đã xây dựng thành công Hệ thống Quản lý Sân bay SkyGate với các kết quả chính:', sa=Pt(6))
for t in [
    'Hệ thống ~15 lớp chính với 3 phân cấp kế thừa (Person → Passenger/Staff → Pilot/GroundStaff; '
    'Aircraft → 3 lớp con; User → Admin/Staff/Customer), quan hệ composition và aggregation.',
    'Áp dụng đầy đủ 4 tính chất OOP: đóng gói, kế thừa, đa hình (dynamic dispatch), trừu tượng '
    '(4 lớp abstract: Person, Staff, Aircraft, User).',
    'Sử dụng 2 Design Pattern: Factory Pattern (AircraftFactory, makeUser) và Composition.',
    'State Machine cho chuyến bay với 12 trạng thái và điều kiện chuyển trạng thái nghiêm ngặt.',
    'Hai giao diện (CLI và Web) dùng chung một AirportSystem backend — tách biệt business logic '
    'và presentation.',
    'Hệ thống phân quyền 3 vai trò với kiểm tra server-side; menu() đa hình quyết định giao diện.',
    'Xử lý 16+ ràng buộc nghiệp vụ phức tạp (chứng chỉ, giờ bay, nghỉ, tương thích hạ tầng, xung '
    'đột lịch, hành lý, trạng thái chuyến, No-show).',
    'Hệ thống đặt/huỷ vé hoàn chỉnh với sơ đồ chọn ghế và phân hạng vé.',
    'Lưu trữ bằng CSDL SQLite chế độ WAL — CLI và Web truy cập đồng thời, dữ liệu nhất quán; tự '
    'động di trú từ định dạng .txt cũ.',
]:
    bullet(t)

h2('6.2. Hạn chế')
for t in [
    'Chưa có unit test tự động — kiểm thử thủ công qua CLI driver và giao diện Web.',
    'Mô phỏng thời tiết xấu còn đơn giản — chỉ hoãn/huỷ thủ công, chưa theo thời gian thực.',
    'Giao diện CLI và Web chưa đồng bộ 100% chức năng (CLI có menu điều phối/thời tiết riêng).',
    'Chưa hỗ trợ đa ngôn ngữ — chuỗi hiển thị là tiếng Việt cố định trong code.',
    'Web client cập nhật bằng cách poll GET /api/state, chưa có push real-time (WebSocket/SSE).',
    'Mật khẩu tài khoản lưu dạng plaintext (phù hợp bài tập, không dùng cho sản phẩm thật).',
]:
    bullet(t)

# ============================================================

# ============================================================
# PHỤ LỤC
# ============================================================
d.add_page_break()
h1('PHẦN PHỤ LỤC')

h2('Tài liệu tham khảo')
for t in [
    'Deitel, P., & Deitel, H. (2017). C++ How to Program (10th ed.). Pearson.',
    'Stroustrup, B. (2013). The C++ Programming Language (4th ed.). Addison-Wesley.',
    'Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (1994). Design Patterns. Addison-Wesley.',
    'cppreference.com — C++ và C++17 Standard Library reference. https://en.cppreference.com/',
    'yhirose/cpp-httplib — C++ header-only HTTP server. https://github.com/yhirose/cpp-httplib',
    'SQLite — Documentation & Amalgamation. https://www.sqlite.org/',
    'Mermaid.js — Diagramming tool for UML. https://mermaid.js.org/',
    'ISO/IEC 14882:2017 — Programming languages — C++ (C++17 standard).',
]:
    bullet(t)

h2('Phụ lục A: Mã nguồn các lớp chính')
para('Toàn bộ mã nguồn được tổ chức theo module trong thư mục src/ (aircraft/, people/, '
     'operations/, system/, auth/, common/, web/). Các đoạn code minh họa cho nguyên lý OOP đã '
     'trích dẫn trong Chương 4. Mã nguồn đầy đủ xem tại repository GitHub (Phụ lục B).', sa=Pt(6))

h2('Phụ lục B: Link GitHub Repository & Demo trực tuyến')
para('Repository GitHub: https://github.com/phhuy21/OOP', sb=Pt(6))
para('Link trải nghiệm trực tuyến: http://159.65.141.209:8888')
para('Tài khoản demo: admin/admin123 (Admin) · staff/staff123 (Staff) · customer/cus123 (Customer).')
para('Repository bao gồm: mã nguồn C++ (src/), giao diện Web (web/), thư viện ngoài (third_party/: '
     'httplib.h, sqlite3.c/.h), script build (build.sh, build_web.sh), tài liệu README.md và CLAUDE.md.',
     italic=True)

# ============================================================
d.save('BAO_CAO_SKYGATE.docx')
print('SUCCESS: Saved BAO_CAO_SKYGATE.docx')
